"""
File processing service for QuizSense.
Handles text extraction from PDF and Word files using:
  - PyMuPDF (fitz) for standard PDFs (primary, fastest)
  - PyPDF2 as fallback text extraction
  - AI Vision API for cloud-based OCR (scanned PDFs)
  - python-docx for Word documents
"""

import base64
import io
import logging
import zipfile

import docx
import fitz  # PyMuPDF
import PyPDF2

logger = logging.getLogger(__name__)

PAGE_TEXT_MIN_CHARS = 50
AI_VISION_OCR_MAX_PAGES = 100


def extract_text_from_pdf(file_obj):
    """
    Extract text from a PDF file.
    Processes each page independently: PyMuPDF first, PyPDF2 fallback, then
    AI Vision OCR only for low-text/scanned pages up to the configured cap.
    """
    from django.conf import settings

    page_texts = []
    ocr_pages_used = 0
    ocr_max_pages = getattr(settings, "AI_VISION_OCR_MAX_PAGES", AI_VISION_OCR_MAX_PAGES)

    try:
        file_obj.seek(0)
        file_bytes = file_obj.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        logger.warning("PyMuPDF document open failed: %s", e)
        file_obj.seek(0)
        return _parse_pdf_pypdf2(file_obj).strip()

    pypdf_reader = None
    try:
        pypdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        logger.warning("PyPDF2 reader setup failed: %s", e)

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = ""

            try:
                page_text = (page.get_text() or "").strip()
            except Exception as e:
                logger.warning("PyMuPDF extraction failed on page %d: %s", page_num + 1, e)

            if len(page_text) < PAGE_TEXT_MIN_CHARS and pypdf_reader is not None:
                try:
                    fallback_text = (pypdf_reader.pages[page_num].extract_text() or "").strip()
                    if len(fallback_text) > len(page_text):
                        page_text = fallback_text
                except Exception as e:
                    logger.warning("PyPDF2 extraction failed on page %d: %s", page_num + 1, e)

            if len(page_text) < PAGE_TEXT_MIN_CHARS and ocr_pages_used < ocr_max_pages:
                ocr_text = _ocr_pdf_page_ai_vision(page, page_num + 1)
                ocr_pages_used += 1
                if len(ocr_text.strip()) > len(page_text):
                    page_text = ocr_text.strip()
            elif len(page_text) < PAGE_TEXT_MIN_CHARS and ocr_pages_used >= ocr_max_pages:
                logger.info(
                    "[AI_VISION_OCR] Page %d skipped because OCR cap reached (%d pages)",
                    page_num + 1,
                    ocr_max_pages,
                )

            if page_text:
                page_texts.append(f"--- Page {page_num + 1} ---\n{page_text}")
    finally:
        doc.close()

    return "\n\n".join(page_texts).strip()


def extract_text_from_docx(file_obj):
    """
    Extract text from a Word (.docx) document.

    Reads normal paragraphs, tables, headers/footers, and raw Word XML text.
    If no real Word text exists, OCR embedded images from image-only DOCX files.
    """
    from django.conf import settings

    file_obj.seek(0)
    file_bytes = file_obj.read()
    text_parts = []

    try:
        document = docx.Document(io.BytesIO(file_bytes))

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        for table in document.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = " ".join(
                        para.text.strip()
                        for para in cell.paragraphs
                        if para.text.strip()
                    )
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        for section in document.sections:
            for header_footer in (section.header, section.footer):
                for para in header_footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        text_parts.append(text)
    except Exception as e:
        logger.warning("DOCX structured extraction failed: %s", e)

    xml_text = _extract_docx_xml_text(file_bytes)
    if xml_text:
        existing = "\n".join(text_parts)
        if len(xml_text) > len(existing):
            text_parts = [xml_text]

    extracted_text = "\n".join(text_parts).strip()
    if extracted_text:
        logger.info("DOCX text extracted (%d chars)", len(extracted_text))
        return extracted_text

    ocr_max_pages = getattr(settings, "AI_VISION_OCR_MAX_PAGES", AI_VISION_OCR_MAX_PAGES)
    logger.info("DOCX returned no Word text; attempting OCR on embedded images (cap=%d)", ocr_max_pages)
    return _ocr_docx_embedded_images(file_bytes, max_images=ocr_max_pages).strip()


def _extract_docx_xml_text(file_bytes):
    """Fallback extractor for Word XML text nodes, including text boxes/content controls."""
    try:
        from xml.etree import ElementTree as ET

        parts = []
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            xml_names = [
                name for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]
            for name in xml_names:
                try:
                    root = ET.fromstring(archive.read(name))
                except Exception:
                    continue
                for node in root.iter():
                    if node.tag.endswith("}t") and node.text and node.text.strip():
                        parts.append(node.text.strip())
        return "\n".join(parts).strip()
    except Exception as e:
        logger.warning("DOCX XML text extraction failed: %s", e)
    return ""


def _ocr_docx_embedded_images(file_bytes, max_images):
    """OCR embedded DOCX images for scanned/image-only Word files."""
    text_parts = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            image_names = [
                name for name in archive.namelist()
                if name.startswith("word/media/")
                and name.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"))
            ]
            image_names.sort()
            for index, name in enumerate(image_names[:max_images], start=1):
                image_bytes = archive.read(name)
                image_text = _ocr_image_bytes_ai_vision(image_bytes, label=f"DOCX image {index}")
                if image_text:
                    text_parts.append(f"--- DOCX Image {index} ---\n{image_text}")

            if len(image_names) > max_images:
                logger.info(
                    "[AI_VISION_OCR] DOCX OCR skipped %d images because cap reached (%d images)",
                    len(image_names) - max_images,
                    max_images,
                )
    except Exception as e:
        logger.warning("DOCX embedded image OCR failed: %s", e)
    return "\n\n".join(text_parts)


def _parse_pdf_pymupdf(file_obj):
    """Fast PDF text extraction via PyMuPDF (fitz)."""
    text_parts = []
    try:
        file_obj.seek(0)
        doc = fitz.open(stream=file_obj.read(), filetype="pdf")
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)
        doc.close()
    except Exception as e:
        logger.warning("PyMuPDF extraction failed: %s", e)
    return "\n".join(text_parts)


def _parse_pdf_pypdf2(file_obj):
    """Fallback PDF text extraction via PyPDF2."""
    text_parts = []
    try:
        file_obj.seek(0)
        reader = PyPDF2.PdfReader(file_obj)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    except Exception as e:
        logger.warning("PyPDF2 extraction failed: %s", e)
    return "\n".join(text_parts)


def _ocr_pdf_page_ai_vision(page, page_number):
    """
    Cloud-based OCR using AI Vision API for a single rendered PDF page.
    """
    try:
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        page_text = _ocr_image_bytes_ai_vision(img_bytes, label=f"page {page_number}")
        if page_text:
            logger.info("[AI_VISION_OCR] Page %d extracted (%d chars)", page_number, len(page_text))
        return page_text
    except Exception as e:
        logger.warning("AI Vision OCR failed on page %d: %s", page_number, e)
    return ""


def _ocr_image_bytes_ai_vision(image_bytes, label):
    """Send image bytes to AI Vision OCR and return extracted text.

    Primary model: Qwen2.5-VL-72B-Instruct (best OCR quality)
    Fallback model: Llama-3.2-90B-Vision-Instruct (retry on failure)
    """
    try:
        import json
        import requests
        from django.conf import settings

        api_key = getattr(settings, 'AI_PROVIDER_API_KEY', '')
        if not api_key:
            logger.warning("AI Vision OCR skipped for %s — AI_PROVIDER_API_KEY not set", label)
            return ""

        img_b64 = base64.b64encode(image_bytes).decode("utf-8")

        OCR_PROMPT = (
            "You are an expert OCR system. Extract ALL text from this image with high accuracy.\n"
            "Rules:\n"
            "- Preserve the original text structure, headings, and paragraphs.\n"
            "- For tables, represent them clearly with rows and columns.\n"
            "- For code or formulas, preserve all syntax and special characters exactly.\n"
            "- If text is blurry or unclear, use context to infer the most likely characters.\n"
            "- Do NOT add any commentary, explanations, or introductory text.\n"
            "- Return ONLY the extracted text."
        )

        models = [
            "Qwen/Qwen2.5-VL-72B-Instruct",
            "meta-llama/Llama-3.2-90B-Vision-Instruct",
        ]

        for model in models:
            try:
                text = _call_vision_model(model, api_key, img_b64, OCR_PROMPT, label)
                if text:
                    return text
            except Exception as e:
                logger.warning("[AI_VISION_OCR] Model %s failed on %s: %s", model, label, e)
                continue

        logger.warning("[AI_VISION_OCR] All vision models failed on %s", label)
        return ""

    except Exception as e:
        logger.warning("AI Vision OCR failed on %s: %s", label, e)
    return ""


def _call_vision_model(model, api_key, img_b64, prompt, label):
    """Call a single vision model and return extracted text."""
    import json
    import requests

    url = "https://api.deepinfra.com/v1/openai/chat/completions"

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt,
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{img_b64}"
                        },
                    },
                ],
            }
        ],
        "max_tokens": 4096,
        "temperature": 0.0,
    }

    response = requests.post(url, headers=headers, json=payload, timeout=90)
    if response.status_code != 200:
        logger.warning("[AI_VISION_OCR] HTTP %d on %s (%s): %s", response.status_code, model, label, response.text[:500])
        return ""

    data = response.json()
    if "choices" not in data:
        logger.warning("[AI_VISION_OCR] Unexpected response on %s (%s): %s", model, label, json.dumps(data)[:500])
        return ""

    text = data["choices"][0]["message"]["content"].strip()
    if text:
        logger.info("[AI_VISION_OCR] %s extracted with %s (%d chars)", label, model, len(text))
    return text
